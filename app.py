import streamlit as st
import google.generativeai as genai
from docx import Document
import io
import os
import time
import re
from supabase import create_client, Client
import pandas as pd

# --- 模型配置区 (锁定版) ---
MODEL_EVAL = "gemini-3.1-pro-preview" 
MODEL_ADAPTIVE = "models/gemini-2.5-flash"
MODEL_CREATIVE = "gemini-2.5-flash"

st.error("🚀 当前代码版本：商业漏斗部署版 (前后台解耦大屏增强)")

# --- 🌟 2. Session State 初始化 ---
init_keys = {
    "user": None, "access_token": None, "refresh_token": None, # 🔑 通行证槽位
    "access_granted": False, "is_vip": False, 
    "is_open_test": False,
    "e_report": None, "e_score": 0, "e_date": "", "e_work_title": "",
    "c_guide": None, "leaderboard": [], "last_eval_time": 0.0,
    "last_creative_prompt": "", "last_eval_text": ""
}
for key, value in init_keys.items():
    if key not in st.session_state:
        st.session_state[key] = value

# --- 专家标注归档逻辑 ---
def save_negative_sample(model_name, prompt, output, critique):
    source_tag = " [CREATIVE]" if "创作" in model_name else " [EVAL]"
    with st.spinner("数据入库中..."):
        try:
            payload = {
                "model_name": model_name + source_tag,
                "prompt": prompt,
                "ai_output": output,
                "expert_critique": critique,
                "user_id": st.session_state['user'].id if st.session_state.get('user') else None
            }
            supabase.table("nal_negative_samples").insert(payload).execute()
            st.toast("✅ 样本已记录，感谢您的专家意见！")
            return True
        except Exception as e:
            st.error(f"同步失败: {e}")
            return False

# --- 顶部标题区 ---
st.set_page_config(page_title="NAL | 新艺文社数字化平台", page_icon="nal_logo.jpg", layout="wide") 

st.markdown("""
    <h1 style='text-align: center; margin-bottom: 0;'>NewArtLiterature Collective (NAL)</h1>
    <h3 style='text-align: center; margin-top: 0; color: #555;'>新艺文社数字化文学平台</h3>
    """, unsafe_allow_html=True)
st.divider()

def clean_text(text):
    if not isinstance(text, str): return ""
    return "".join(c for c in text if c.isprintable() or c in "\n\r\t")

# --- 🔒 3. 环境变量与 API 配置 ---
try:
    api_key = os.getenv("GEMINI_API_KEY")
    if api_key: genai.configure(api_key=api_key)
    else: st.error("❌ 错误：未检测到 GEMINI_API_KEY。"); st.stop()
except Exception as e:
    st.error(f"无法配置 Gemini API: {e}")

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# 🔑 核心记忆恢复：向 Supabase 客户端注入登录凭证
if st.session_state.get('access_token') and st.session_state.get('refresh_token'):
    try:
        supabase.auth.set_session(
            st.session_state['access_token'], 
            st.session_state['refresh_token']
        )
    except Exception:
        pass 

# --- 🗄️ 全局数据拉取函数 ---
@st.cache_data(ttl=3600)
def fetch_nal_models_from_db():
    res = supabase.table("evaluation_models").select("name, parameters, description").execute()
    matrix = {row["name"]: row["parameters"] for row in res.data}
    descriptions = {row["name"]: row.get("description", "暂无简介") for row in res.data}
    return matrix, descriptions

@st.cache_data(ttl=30) 
def fetch_competition_data():
    """分页滚动拉取全量大赛数据,穿透 PostgREST 单次查询 1000 行的默认截断。
    ⚠️ 此前未分页且按 flash_score 降序,被截断的恰好是低分作品,
    导致总收稿量恒为 1000、争议/AI 拦截数被系统性低估。"""
    try:
        all_rows = []
        page_size = 1000
        start = 0
        while True:
            res = supabase.table("nal_archives").select(
                "id, work_title, genre, flash_score, is_controversial, review_status, score, final_award, committee_summary, created_at, ai_risk_score, is_ai_suspected"
            ).eq("archive_type", "competition_2026")\
            .order("flash_score", desc=True)\
            .range(start, start + page_size - 1).execute()
            if not res.data:
                break
            all_rows.extend(res.data)
            if len(res.data) < page_size:
                break
            start += page_size
        return all_rows
    except Exception as e:
        print(f"数据库读取失败: {e}") 
        return []

@st.cache_data(ttl=30)
def fetch_final_report_from_db():
    try:
        res = supabase.table("nal_archives").select("content, created_at").eq("archive_type", "competition_2026_report").order("created_at", desc=True).limit(1).execute()
        return res.data[0] if res.data else None
    except Exception:
        return None

EVAL_SYSTEM_MATRIX, MODEL_DESCRIPTIONS = fetch_nal_models_from_db()

# --- 🗄️ 档案室自动归档核心逻辑 (优雅容错版) ---
def save_to_nal_archive(archive_type, title, content, score=0):
    if st.session_state.get('user'):
        u_id = st.session_state['user'].id
        
        with st.spinner(f"💾 正在同步《{title}》至云端档案室..."):
            try:
                payload = {
                    "user_id": u_id,
                    "archive_type": archive_type,
                    "work_title": title,
                    "content": content,
                    "score": score
                }
                response = supabase.table("nal_archives").insert(payload).execute()
                
                if response.data:
                    st.toast(f"✅ 存档成功：{title}", icon="📁")
                    time.sleep(1.0) 
                    return True
                else:
                    # 🚨 移除 st.stop()，改为温和报错并返回 False
                    st.error(f"🚨 数据库拒绝写入。可能是权限或网络原因，请注意保存本地副本。")
                    return False
            except Exception as e:
                # 🚨 移除 st.stop()，改为温和报错并返回 False
                st.error(f"🚨 档案系统同步物理报错: {e}")
                return False 
    return False
    
def get_eval_model_from_db(model_name):
    try:
        response = supabase.table("evaluation_models").select("system_instruction, parameters, description").eq("name", model_name).single().execute()
        return response.data
    except Exception as e:
        st.error(f"无法调取数据库模型 '{model_name}': {e}")
        return None    

def get_adaptive_instruction(model_data, current_text, user_note=""):
    base_params = model_data.get('parameters', {})
    if not base_params: return model_data.get('system_instruction', '')
    
    sense_prompt = """你是一个文本特征分析器。请严谨分析该儿童文学文本指标(0.0-1.0)：
    1.fantasy(幻想感) 2.reality(现实/时代感) 3.character(人物心理深度)。
    必须仅输出纯 JSON 格式：{"fantasy": 0.5, "reality": 0.5, "character": 0.5}"""

    try:
        feature_model = genai.GenerativeModel(MODEL_ADAPTIVE) 
        f_res = feature_model.generate_content(
            f"{sense_prompt}\n内容：{current_text[:2000]}", 
            generation_config=genai.types.GenerationConfig(response_mime_type="application/json", temperature=0.1)
        )
        import json
        features = json.loads(f_res.text)
    except Exception as e:
        features = {"fantasy": 0.5, "reality": 0.5, "character": 0.5}
    
    adjusted_weights = base_params.copy()
    sensitivity = 15
    mapping = {
        "fantasy": ["跨界", "共鸣", "幻想", "想象", "诗意", "隐喻", "对位", "意象", "视觉", "分镜", "审美", "张力", "形式", "艺术", "留白", "介入", "童话", "超自然", "虚构"],
        "reality": ["时代", "社会", "技术", "异化", "现实", "真相", "背景", "偏见", "价值观", "文化", "伦理", "生态", "教育", "批判", "成人主义", "意识形态", "显性", "潜意识", "病灶"],
        "character": ["人物", "心理", "契合", "塑造", "成长", "主体", "非人类", "尊严", "读者", "共生", "视角", "动机", "弧光", "自我", "生命本位", "空间", "体验", "共情"]
    }

    for dim in adjusted_weights.keys():
        if any(k in dim for k in mapping["fantasy"]): adjusted_weights[dim] = max(1, adjusted_weights[dim] + (features.get('fantasy', 0.5) - 0.5) * sensitivity)
        if any(k in dim for k in mapping["reality"]): adjusted_weights[dim] = max(1, adjusted_weights[dim] + (features.get('reality', 0.5) - 0.5) * sensitivity)
        if any(k in dim for k in mapping["character"]): adjusted_weights[dim] = max(1, adjusted_weights[dim] + (features.get('character', 0.5) - 0.5) * sensitivity)

    intervention_log = ""
    if user_note:
        for dim in adjusted_weights.keys():
            if dim[:2] in user_note:
                adjusted_weights[dim] += 25
                intervention_log += f"【已根据备注强化‘{dim}’】 "

    total = sum(adjusted_weights.values())
    final_weights = {k: round((v/total)*100, 1) for k, v in adjusted_weights.items()}
    weight_desc = "\n".join([f"- {k}: {v}%" for k, v in final_weights.items()])
    return f"""{model_data['system_instruction']}\n---\n【NAL 通用自适应校准报告】\n文本指纹：幻想({features.get('fantasy')})，现实({features.get('reality')})，人物({features.get('character')})\n{intervention_log}\n动态权重矩阵：\n{weight_desc}\n---\n请按此分配执行评审。"""

MODEL_OPTIONS = [
    "全景综合-通用基准模型", "NAL-首席专家锐评模型", "李利芳-儿童文学价值模型",
    "朱自强-儿童本位论模型", "视觉叙事-图文对位模型", "霍林代尔-意识形态批判模型", "后人类/生态主义先锋模型"
]

# --- 🌟 5. 路由逻辑与支付墙拦截 ---
is_internal_mode = st.query_params.get("mode") == "internal"
is_saas_mode = not is_internal_mode

if is_saas_mode:
    if st.session_state['user'] is None:
        st.title("🌟 NAL 商业版 (SaaS)")
        c1, c2 = st.columns(2)
        with c1:
            st.subheader("🔑 现有会员登录")
            l_e = st.text_input("邮箱")
            l_p = st.text_input("密码", type="password")
            if st.button("立即登录"): 
                try:
                    res = supabase.auth.sign_in_with_password({"email": l_e, "password": l_p})
                    st.session_state['user'] = res.user
                    st.session_state['access_token'] = res.session.access_token
                    st.session_state['refresh_token'] = res.session.refresh_token
                    st.rerun()
                except Exception as e: st.error(f"登录失败，原因: {e}")
        with c2:
            st.subheader("📝 注册新会员")
            r_e = st.text_input("常用邮箱")
            r_p = st.text_input("设置密码")
            if st.button("免费注册"):
                try:
                    supabase.auth.sign_up({"email": r_e, "password": r_p})
                    st.success("注册成功！请检查邮箱或直接登录。")
                except Exception as e: st.error(f"注册错误: {e}")
        st.stop()
    else:
        if not st.session_state['is_vip']:
            st.title("🚀 解锁 NAL 专业版双擎系统")
            st.markdown(f"欢迎您，**{st.session_state['user'].email}**！您目前使用的是未激活账户，请订阅以获取 AI 引擎的完整算力支持。")
            c_pay1, c_pay2, c_pay3 = st.columns([1, 2, 1])
            with c_pay2:
                st.info("💎 **NAL Pro 创作者订阅**\n\n- 开启无限次创作推演\n- 启动极其严苛的 智能专家 评审")
                if st.button("🛠️ [开发者通道] 模拟支付成功，一键激活", width="stretch"):
                    st.session_state['is_vip'] = True
                    st.rerun()
            st.stop()
else:
    if not st.session_state["access_granted"]:
        st.title("🔒 NAL 内部测试系统")
        inv = st.text_input("评委/作者邀请码：", type="password")
        if st.button("确认进入"): 
            env_codes = os.getenv("NAL_INVITE_CODES")
            VALID_CODES = [code.strip() for code in env_codes.split(",")] if env_codes else ["NAL2026"]
            
            if inv == "Open_test":
                st.session_state["access_granted"] = True
                st.session_state["is_open_test"] = True  
                st.rerun()
            elif inv in VALID_CODES: 
                st.session_state["access_granted"] = True
                st.session_state["is_open_test"] = False 
                st.rerun()
            else: 
                st.error("邀请码无效。")
        st.stop()

# --- 唯一的侧边栏控制块 ---
with st.sidebar:
    enable_critique = False
    st.image("nal_logo.jpg", width="stretch")
    st.markdown("<h3 style='text-align: center;'>NAL 控制台</h3>", unsafe_allow_html=True)
    
    if is_saas_mode:
        if st.session_state.get('user'):
            st.success(f"已登录: {st.session_state['user'].email}")
            enable_critique = st.toggle("🎨 开启专家纠偏模式", value=False, key="saas_critique_toggle")
            if st.button("🚪 退出登录", width="stretch"):
                supabase.auth.sign_out()
                st.session_state['user'] = None
                st.session_state['access_token'] = None  
                st.session_state['refresh_token'] = None 
                st.session_state['is_vip'] = False
                st.rerun()
    else:
        is_open_test = st.session_state.get("is_open_test", False)
        if is_open_test:
            st.warning("模式：开放测试（自动反馈激活）")
            enable_critique = True
        else:
            st.info("模式：正式成员")
            enable_critique = st.toggle("🎨 开启专家纠偏模式", value=False)

    st.divider()

    if enable_critique:
        last_action = st.session_state.get('last_action')
        options = []
        if last_action == "eval" and st.session_state.get('e_report'): options = ["评审报告"]
        elif last_action == "creative":
            if st.session_state.get('last_outline'): options.append("创作大纲")
            if st.session_state.get('last_snippet'): options.append("创作片段")

        if options:
            st.markdown("#### 🛠️ 样本纠偏")
            target = st.radio("选择对象：", options, key="side_target_radio")
            if target == "评审报告":
                curr_out = st.session_state.get('e_report')
                curr_model = st.session_state.get("last_eval_model", "Unknown")
                curr_prompt = st.session_state.get("last_eval_text", "")
            else:
                curr_out = st.session_state.get('last_outline') if target == "创作大纲" else st.session_state.get('last_snippet')
                curr_model = f"{st.session_state.get('last_mentor_used')}-{target}"
                curr_prompt = st.session_state.get("last_creative_prompt", "")

            expert_thought = st.text_area("纠偏意见：", key="side_critique_input", placeholder="输入您的想法...")
            if st.button("🚨 提交样本", width="stretch", type="primary"):
                if expert_thought:
                    save_negative_sample(curr_model, curr_prompt, curr_out, expert_thought)
                    st.rerun()
        else:
            st.caption("👈 完成生成后此处开启纠偏")

# --- 🌟 6. 主功能界面 ---
st.title("NAL 数字化文学双擎系统")

tabs_list = ["💡 创作伴侣", "⚖️ 深度评审", "🏆 评审排行榜", "📁 我的档案室"]

is_logged_in_saas = is_saas_mode and st.session_state.get('user') is not None
if is_logged_in_saas:
    tabs_list.append("🏆 大赛监控中心")

tabs = st.tabs(tabs_list)
tab1, tab2, tab3, tab4 = tabs[0], tabs[1], tabs[2], tabs[3]

with tab1:
    st.header("💡 NAL/学术体系 智能创作指导系统")
    mentor_options = list(MODEL_DESCRIPTIONS.keys())
    mentor_type = st.selectbox("💡 请选择您的创作指导导师体系：", MODEL_OPTIONS, key="mentor_select")
    st.info(f"**导师风格**：{MODEL_DESCRIPTIONS.get(mentor_type, '')}")
            
    u_prompt = st.text_area("输入您的灵感片段：", placeholder="输入主题、核心冲突或想要探讨的时代命题...", height=150, key="u_input_2026")
    c_filename = st.text_input("📄 设定片段导出文件名", value="NAL_Highlights", key="f_input_2026")

    btn_disabled = not u_prompt or u_prompt == st.session_state["last_creative_prompt"]
    if st.button("启动 创作推演", disabled=btn_disabled, key="run_creative_2026"): 
        with st.spinner(f"正在调动【{mentor_type}】进行深度构思与 800 字试写..."):
            try:
                mentor_desc = MODEL_DESCRIPTIONS[mentor_type]
                focus_dimensions = "、".join(EVAL_SYSTEM_MATRIX.get(mentor_type, {}).keys())
                expert_standard = f"【核心指导思想】：{mentor_desc}\n【重点发力维度】：你的大纲和试写必须极力展现以下学术特质：{focus_dimensions}。"

                creative_sys_inst = f"""你现在是儿童文学领域的金牌创作指导。任务：协助作者构思并实打实撰写长片段。\n标准：{expert_standard}\n【输出格式要求】\n1. 前半部分：包含【核心立意升华】、【人物弧光设定】、【情节大纲建议】。\n2. 中间必须插入一行：===片段分割线===\n3. 后半部分：【高光片段试写】。\n注意：试写片段必须达到 600-800 字，要求极强的画面感和文学性，绝不准敷衍！"""
                model = genai.GenerativeModel(model_name=MODEL_CREATIVE, system_instruction=creative_sys_inst)
                
                res = model.generate_content(u_prompt, generation_config=genai.types.GenerationConfig(temperature=0.7, max_output_tokens=8192, top_p=0.95))
                if res.candidates and res.candidates[0].content.parts:
                    full_content = res.text
                    st.session_state['c_guide'] = full_content
                    st.session_state['last_action'] = "creative"
                    st.session_state['last_mentor_used'] = mentor_type 
                    st.session_state['last_creative_prompt'] = u_prompt
                    st.session_state['e_report'] = None 
                
                    parts = re.split(r'创作片段|【创作片段】|===片段分割线===|试写片段', full_content)
                    st.session_state['last_outline'] = parts[0].strip()
                    if len(parts) > 1: st.session_state['last_snippet'] = parts[1].strip()
                    else: st.session_state['last_snippet'] = "（引擎未产出独立片段，请在大纲中查看立意）"
                
                    is_saved = save_to_nal_archive("creative", c_filename, full_content)
                    if is_saved:
                        st.success(f"✅ 创作完成！档案已存至：{c_filename}")
                        st.rerun()
                    else:
                        # 🚨 优雅容错：不刷新页面，保留警告，供用户手动导出文档
                        st.warning("⚠️ 内容已成功生成，但未能同步至云端档案室。")
                else:
                    reason = res.candidates[0].finish_reason if res.candidates else "未知"
                    st.error(f"⚠️ 创作引擎由于安全策略未生成内容 (原因代码: {reason})。请尝试修改关键词或降低敏感度。")
            except Exception as e:
                st.error(f"❌ 创作引擎异常: {e}")
    
    if st.session_state.get('c_guide'):
        guide_text = st.session_state.get('c_guide')
        outline_content, snippet_content = "", ""
        if "===片段分割线===" in guide_text:
            parts = guide_text.split("===片段分割线===")
            outline_content = parts[0].strip()
            snippet_content = parts[1].strip() if len(parts) > 1 else ""
        if not snippet_content or len(snippet_content) < 50:
            fallback_parts = re.split(r'【高光片段试写】', guide_text)
            if len(fallback_parts) > 1:
                outline_content = fallback_parts[0].strip()
                snippet_content = "【高光片段试写】\n" + fallback_parts[1].strip()
        if not outline_content:
            outline_content = guide_text
            snippet_content = "（系统未能自动切分片段，请在上方完整内容中查看）"

        st.divider()
        col_show1, col_show2 = st.columns([2, 3])
        with col_show1:
            st.markdown("#### 📌 创作策划大纲")
            st.info(outline_content)
            d1 = Document(); d1.add_heading('NAL 创作大纲', 0); d1.add_paragraph(clean_text(outline_content))
            b1 = io.BytesIO(); d1.save(b1)
            st.download_button("📥 导出大纲 (.docx)", b1.getvalue(), f"Outline_{c_filename}.docx", key="dl_o_2026")

        with col_show2:
            st.markdown("#### 📖 高光片段试写")
            st.success(snippet_content)
            d2 = Document(); d2.add_heading('NAL 高光片段', 0); d2.add_paragraph(clean_text(snippet_content))
            b2 = io.BytesIO(); d2.save(b2)
            st.download_button("📥 导出片段 (.docx)", b2.getvalue(), f"{c_filename}.docx", key="dl_s_2026")
        
with tab2:
    st.header("⚖️ 深度评审系统")
    curr_t = time.time()
    last_model = st.session_state.get("last_eval_model", "")
    cd_time = 15 - (curr_t - st.session_state.get("last_eval_time", 0))
    selected_model = st.selectbox("⚖️ 请选择评审学术体系：", MODEL_OPTIONS, key="eval_select")
    st.info(f"**学术底色**：{MODEL_DESCRIPTIONS.get(selected_model, '暂无描述')}")
    base_weights = EVAL_SYSTEM_MATRIX.get(selected_model, {})
    
    up_file = st.file_uploader("上传作品文本 (.docx)", type=["docx"], key="eval_uploader")
    eval_note = st.text_area("在此输入评委备注（系统将根据语义自动调整权重）：", placeholder="例如：请重点关注‘时代性’...", height=100)
    
    raw_text = ""
    if up_file:
        try:
            up_file.seek(0)
            doc = Document(up_file)
            raw_text = "\n".join([p.text for p in doc.paragraphs])
        except Exception as e: st.error(f"文件读取失败: {e}")

    text_is_new = raw_text != st.session_state.get("last_eval_text", "")
    model_is_new = selected_model != last_model
    can_proceed = up_file is not None and (text_is_new or model_is_new)
    
    if st.button("启动 智能 评审", disabled=not can_proceed):
        if cd_time > 0:
            st.warning(f"系统冷却中: {int(cd_time)}秒，请稍后再试。")
        elif raw_text:
            with st.spinner(f"正在调取【{selected_model}】并启动自适应引擎..."):
                try:
                    m_data = get_eval_model_from_db(selected_model)
                    if m_data:
                        example_dims = ""
                        for k, v in base_weights.items():
                            example_dims += f"* **{k}**：{int(v * 0.8)}/{v}分 - 这里的描写非常生动...\n"

                        eval_sys_inst = f"""你现在是 NAL 数字化平台的顶级学术评审专家。当前执行体系：【{selected_model}】。维度的【最高满分】分别是：{base_weights}。必须输出真实的个位数字分数！绝不允许打0分！
                        【第一阶段：前置硬伤排查】1.逻辑与事实。2.原创性。
                        【评审准则】1.默认找瑕疵，平庸者60-65。2.落俗套、说教直接扣50%。3.不原创最高70分。
                        【强制输出规范】
                        ### 💡 逻辑与原创性审查\n* **事实与逻辑排查**：[评语]\n* **原创性评估**：[评语]
                        ### 🧮 维度解析与单项得分\n{example_dims}
                        ### 📝 核心修改建议\n1. [建议1]
                        ---\n### 📊 综合学术评分：[总分]/100"""

                        final_inst = get_adaptive_instruction(m_data, raw_text, eval_note)
                        combined_system_instruction = final_inst + "\n\n" + eval_sys_inst

                        current_engine = "gemini-2.5-flash" if st.session_state.get("is_open_test") else MODEL_EVAL
                        eval_model = genai.GenerativeModel(model_name=current_engine, system_instruction=combined_system_instruction)
                        prompt = f"【内容】：\n{raw_text}\n\n【备注】：{eval_note if eval_note else '无'}\n\n请严格按格式打分！"
                        
                        try:
                            res_obj = eval_model.generate_content(prompt, generation_config=genai.types.GenerationConfig(temperature=0.4))
                            if res_obj.candidates and res_obj.candidates[0].content.parts:
                                st.session_state['e_report'] = res_obj.text
                                st.session_state["last_eval_time"] = time.time()
                                st.session_state["last_eval_text"] = raw_text
                                st.session_state["last_eval_model"] = selected_model
                                st.session_state["e_date"] = time.strftime('%Y-%m-%d %H:%M:%S')
                                st.session_state["e_work_title"] = up_file.name.rsplit('.', 1)[0]
                                st.session_state['last_action'] = "eval"
                                st.session_state['last_outline'] = None
                                st.session_state['last_snippet'] = None                       
                                
                                score_val = 0
                                m_score = re.search(r"(?:综合学术评分|综合评分|总分)[】\]]?[:：]?\[?(\d{1,3})\]?(?:/100|分)?", res_obj.text.replace('*', '').replace(' ', ''))
                                if m_score: score_val = int(m_score.group(1))
                                st.session_state['e_score'] = score_val
                                
                                if "leaderboard" not in st.session_state: st.session_state["leaderboard"] = []
                                st.session_state['leaderboard'].append({"作品": up_file.name, "分数": score_val, "日期": st.session_state["e_date"], "体系": selected_model})
                                
                                is_saved = save_to_nal_archive("evaluation", up_file.name, res_obj.text, score_val)
                                if is_saved:
                                    st.rerun()
                                else:
                                    # 🚨 优雅容错：不刷新页面，给用户保留看报告和导出 Word 的机会
                                    st.warning("⚠️ 报告已生成，但因数据库连接问题未能归档。您可以直接在下方查阅并下载报告。")
                            else: st.error("⚖️ 报告生成失败。")
                        except Exception as e: st.error(f"评审发生异常: {e}")
                    else: st.error("未找到模型数据。")
                except Exception as e: st.error(f"评审异常: {e}")

    if st.session_state.get('e_report'):
        r_title = st.session_state.get("last_eval_model", selected_model)
        r_work = st.session_state.get('e_work_title', '未知')
        
        col_m1, col_m2 = st.columns([1, 3])
        with col_m1: st.metric(f"{r_title} 评分", f"{st.session_state.get('e_score', 0)} / 100")
        with col_m2: st.caption(f"📅 {st.session_state.get('e_date')} | 📁 {r_work}")
        st.divider()
        
        full_rpt = st.session_state['e_report']
        display_content = full_rpt.split("---")[-1].strip() if "---" in full_rpt else full_rpt
        if len(display_content) < 100: display_content = full_rpt
        st.markdown(display_content)
        
        rd_doc = Document(); rd_doc.add_heading(f'{r_title} 评审报告', 0); rd_doc.add_paragraph(f"作品：{r_work}\n日期：{st.session_state.get('e_date')}\n")
        try: rd_doc.add_paragraph(clean_text(display_content))
        except NameError: rd_doc.add_paragraph(display_content) 
            
        rb = io.BytesIO(); rd_doc.save(rb)
        safe_r_work = re.sub(r'[\\/*?:"<>|]', "", r_work)
        st.download_button(label=f"📥 下载《{safe_r_work}》评审报告", data=rb.getvalue(), file_name=f"NAL_Report_{safe_r_work}.docx", width="stretch")

with tab3:
    st.header("🏆 NAL 评审作品排行榜")
    if st.session_state['leaderboard']:
        st.table(sorted(st.session_state['leaderboard'], key=lambda x: x['分数'], reverse=True))

with tab4:
    st.header("📁 NAL 云端档案库")
    user_obj = st.session_state.get('user')
    if user_obj:
        try:
            res_db = supabase.table("nal_archives").select("*").eq("user_id", user_obj.id)\
                .neq("archive_type", "competition_2026")\
                .neq("archive_type", "competition_2026_report")\
                .order("created_at", desc=True).execute()
                
            if res_db.data:
                for arc in res_db.data:
                    col_main, col_del = st.columns([0.85, 0.15])
                    with col_main: exp = st.expander(f"【{arc['archive_type']}】{arc['work_title']} - {arc['created_at'][:10]}")
                    with col_del:
                        if st.button("🗑️", key=f"fast_del_{arc['id']}", type="primary", help="永久删除"):
                            supabase.table("nal_archives").delete().eq("id", arc['id']).execute()
                            st.toast(f"已移除：{arc['work_title']}")
                            time.sleep(0.5); st.rerun()

                    with exp:
                        st.write(arc.get('content', '无内容'))
                        st.divider()
                        col_action, col_danger = st.columns([0.7, 0.3])
                        with col_action:
                            doc = Document(); doc.add_paragraph(clean_text(arc.get('content', ''))); bio = io.BytesIO(); doc.save(bio)
                            st.download_button("📥 重新导出 Docx", bio.getvalue(), f"Arc_{arc['id'][:4]}.docx", key=f"dl_{arc['id']}")
                        with col_danger:
                            if st.button("🚨 确认永久删除", key=f"final_del_{arc['id']}", type="primary"):
                                supabase.table("nal_archives").delete().eq("id", arc['id']).execute()
                                st.success("已永久移除"); time.sleep(0.5); st.rerun()
            else:
                st.info("📂 云端档案室暂无记录。 （注：大赛跑批数据及公文均属后台机房专属，不在此处显示）")
        except Exception as e:
            st.error(f"档案加载失败: {e}")
    else:
        st.warning("💡 档案库管理权限仅对注册会员开放。当前模式仅支持浏览本次会话历史。")
        if st.session_state.get('leaderboard'):
            st.markdown("### 🕒 本次会话评审历史")
            st.table(st.session_state['leaderboard'])

# 🚨 【渲染具备官方公文与漏斗数据双重架构的 Tab 5】
if is_logged_in_saas:
    with tabs[4]: 
        col1, col2 = st.columns([4, 1])
        with col1:
            st.header("🏆 2026 全球儿童文学大赛 - 漏斗式监控大屏")
            st.info("实时同步 NAL 核心机房三段式评审流：Flash 极速海选 ➡️ 大盘配额统筹 ➡️ Pro 铁三角终评。")
        with col2:
            st.write("") 
            if st.button("🔄 刷新实时进度", width="stretch"):
                fetch_competition_data.clear() 
                fetch_final_report_from_db.clear()
                st.rerun() 
        
        official_report = fetch_final_report_from_db()
        if official_report:
            st.subheader("📜 赛事组委会官方决议报告")
            with st.container(border=True):
                st.markdown(official_report['content'])
                st.write("")
                
                doc_rep = Document()
                doc_rep.add_heading('NAL 2026 大赛终评官方决议报告', 0)
                doc_rep.add_paragraph(clean_text(official_report['content']))
                bio_rep = io.BytesIO(); doc_rep.save(bio_rep)
                
                st.download_button(
                    label="🖨️ 下载并打印官方决议公文 (.docx)",
                    data=bio_rep.getvalue(),
                    file_name="NAL_2026_Official_Resolution.docx",
                    type="secondary",
                    width="stretch"
                )
            st.write("")

        comp_data = fetch_competition_data()
        
        if not comp_data:
            st.info("📭 目前数据库中还没有大赛数据。请先在本地终端运行 python genre_ensemble_runner.py ！")
        else:
            df_comp = pd.DataFrame(comp_data)
            
            total_entries = len(df_comp)
            shortlisted_df = df_comp[df_comp['review_status'].isin(['shortlisted', 'pro_done', 'pro_reviewing'])]
            pro_done_df = df_comp[df_comp['review_status'] == 'pro_done']
            controversial_df = df_comp[df_comp['is_controversial'] == True]
            ai_suspect_df = df_comp[df_comp['is_ai_suspected'] == True]
            
            st.markdown("### 📊 评审漏斗实时数据")
            col_c1, col_c2, col_c3, col_c4, col_c5 = st.columns(5)
            col_c1.metric("📥 总收稿量 (海选池)", f"{total_entries} 份")
            col_c2.metric("🎯 成功入围 (5% 配额)", f"{len(shortlisted_df)} 份")
            col_c3.metric("🏆 完成终评 (Pro矩阵)", f"{len(pro_done_df)} 份")
            col_c4.metric("⚠️ 争议打捞 (需人工)", f"{len(controversial_df)} 份")
            col_c5.metric("🤖 AI 嫌疑拦截", f"{len(ai_suspect_df)} 份")
            
            st.divider()

            st.subheader("🌊 赛事全景与 AIGC 拦截公示板")
            
            # 填补空值，防旧数据报错
            df_comp['final_award'] = df_comp['final_award'].fillna('暂未获奖')
            df_comp['score'] = df_comp['score'].fillna(0).astype(int)
            df_comp['ai_risk_score'] = df_comp['ai_risk_score'].fillna(0).astype(int)
            df_comp['is_ai_suspected'] = df_comp['is_ai_suspected'].fillna(False) 
            
            # 生成 AI 警告列 (已提前规避 KeyError 漏洞)
            df_comp['AI_警告'] = df_comp.apply(lambda x: f"🚨 {x['ai_risk_score']}%" if x['is_ai_suspected'] else "✅ 纯天然", axis=1)
            
            # 字段重命名提取
            show_df1 = df_comp[['work_title', 'genre', 'AI_警告', 'flash_score', 'score', 'final_award', 'review_status']].rename(
                columns={
                    'work_title': '作品名称', 
                    'genre': '赛道体裁', 
                    'AI_警告': 'AIGC 风险',
                    'flash_score': '海选基础分', 
                    'score': '终评分',
                    'final_award': '🏅 授予荣誉',
                    'review_status': '流转状态'
                }
            )
            
            status_map = {'flash_done': '待统筹配额', 'rejected': '止步初选', 'shortlisted': '入围待终评', 'pro_reviewing': '终审锁定中', 'pro_done': '终评已决议'}
            show_df1['流转状态'] = show_df1['流转状态'].map(status_map).fillna(show_df1['流转状态'])
            
            # 摒弃中文 Unicode 排序陷阱，直接按真实数值分数降序排序
            show_df1 = show_df1.sort_values(by=['终评分', '海选基础分'], ascending=[False, False])
            
            st.dataframe(show_df1, width="stretch", hide_index=True)
            
            st.write("---") 
            
            st.subheader("📜 最终阶段：终评决议详情公示")
            if pro_done_df.empty:
                st.caption("⏳ 终评池暂无数据，等待引擎进入第三阶段...")
            else:
                for index, row in pro_done_df.iterrows():
                    tier = row['final_award'] if pd.notna(row.get('final_award')) else "未定级"
                    card_title = f"📖 《{row['work_title']}》 ｜ 🏆 终评分：{row.get('score', 0)} 分 ｜ 🏷️ 评级：{tier}"
                    
                    with st.expander(card_title):
                        st.write("")
                        st.markdown("**⚖️ 评审委员会主席最终签署决议：**")
                        if pd.notna(row.get('committee_summary')) and row['committee_summary']:
                            st.info(row['committee_summary'])
                        else:
                            st.warning("⚠️ 监测到该条历史数据尚未生成决议。")
            
            st.divider()
            
            st.subheader("📥 大赛实时台账下载")
            csv_data = show_df1.to_csv(index=False).encode('utf-8-sig')
            st.download_button(
                label="📊 一键导出当前【海选与入围全景状态】报表 (CSV)", 
                data=csv_data, 
                file_name="2026_NAL_Funnel_Status.csv",
                type="primary",
                width="stretch"
            )
